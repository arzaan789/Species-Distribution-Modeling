"""Build editable Ecological Modelling submission documents from reviewed sources."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
MANUSCRIPT = ROOT / "manuscript"
FIGURES = MANUSCRIPT / "figures"

FIGURE_FILES = {
    "Figure 1.": "simulation_workflow.png",
    "Figure 2.": "paired_truth_contrasts.png",
    "Figure 3.": "source_composition_mechanism.png",
    "Figure 4.": "latent_mixture_contrasts.png",
    "Figure 5.": "empirical_source_contrasts.png",
    "Figure 6.": "empirical_map_contrast.png",
}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=80, bottom=80, end=80) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instruction, end])


def configure_document(document: Document, *, letter: bool = False) -> None:
    section = document.sections[0]
    if not letter:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
    section.top_margin = Cm(2.35)
    section.bottom_margin = Cm(2.35)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after = Pt(6)

    for name, size in (("Title", 17), ("Heading 1", 14), ("Heading 2", 12), ("Heading 3", 11)):
        style = document.styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
    document.styles["Title"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_ppr = document.styles["Title"]._element.get_or_add_pPr()
    title_border = title_ppr.find(qn("w:pBdr"))
    if title_border is not None:
        title_ppr.remove(title_border)

    if "Caption" in document.styles:
        caption = document.styles["Caption"]
    else:
        caption = document.styles.add_style("Caption", WD_STYLE_TYPE.PARAGRAPH)
    caption.font.name = "Times New Roman"
    caption.font.size = Pt(9)
    caption.font.color.rgb = RGBColor(0, 0, 0)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    caption.paragraph_format.line_spacing = 1.0
    caption.paragraph_format.space_after = Pt(6)
    caption.paragraph_format.keep_with_next = True

    footer = section.footer.paragraphs[0]
    add_page_number(footer)
    footer.runs[0].font.name = "Arial"
    footer.runs[0].font.size = Pt(9)


def add_inline_markdown(paragraph, text: str) -> None:
    """Add a small, controlled subset of Markdown emphasis."""
    token = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")
    cursor = 0
    for match in token.finditer(text):
        if match.start() > cursor:
            paragraph.add_run(text[cursor : match.start()])
        value = match.group(0)
        if value.startswith("**"):
            run = paragraph.add_run(value[2:-2])
            run.bold = True
        elif value.startswith("*"):
            run = paragraph.add_run(value[1:-1])
            run.italic = True
        else:
            run = paragraph.add_run(value[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        cursor = match.end()
    if cursor < len(text):
        paragraph.add_run(text[cursor:])


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        cells = [cell.strip() for cell in lines[index].strip().strip("|").split("|")]
        rows.append(cells)
        index += 1
    if len(rows) >= 2 and all(re.fullmatch(r":?-{3,}:?", cell) for cell in rows[1]):
        del rows[1]
    return rows, index


def add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    columns = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=columns)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True
    for row_index, values in enumerate(rows):
        if row_index == 0:
            row_properties = table.rows[row_index]._tr.get_or_add_trPr()
            repeat = OxmlElement("w:tblHeader")
            repeat.set(qn("w:val"), "true")
            row_properties.append(repeat)
        for column_index in range(columns):
            cell = table.cell(row_index, column_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            value = values[column_index] if column_index < len(values) else ""
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if column_index < 2 else WD_ALIGN_PARAGRAPH.RIGHT
            paragraph.paragraph_format.line_spacing = 1.0
            paragraph.paragraph_format.space_after = Pt(0)
            add_inline_markdown(paragraph, value)
            for run in paragraph.runs:
                run.font.name = "Arial"
                run.font.size = Pt(8)
                if row_index == 0:
                    run.bold = True
            if row_index == 0:
                set_cell_shading(cell, "E6E6E6")
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_markdown(
    document: Document,
    source: Path | None = None,
    *,
    lines: list[str] | None = None,
    include_figures: bool = False,
) -> None:
    if lines is None:
        if source is None:
            raise ValueError("source or lines is required")
        lines = source.read_text(encoding="utf-8").splitlines()
    index = 0
    equation = False
    while index < len(lines):
        raw = lines[index]
        line = raw.strip()
        if not line:
            index += 1
            continue
        if line == "\\[":
            equation = True
            index += 1
            continue
        if line == "\\]":
            equation = False
            index += 1
            continue
        if equation:
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.keep_together = True
            run = paragraph.add_run(line.replace("\\frac", "frac"))
            run.font.name = "Cambria Math"
            run.font.size = Pt(10)
            run.italic = True
            index += 1
            continue
        if line.startswith("|"):
            rows, index = parse_table(lines, index)
            add_table(document, rows)
            continue
        if line.startswith("# "):
            # The main title is styled on a dedicated title page by the caller.
            if document.paragraphs:
                document.add_heading(line[2:], level=1)
            else:
                document.add_paragraph(line[2:], style="Title")
            index += 1
            continue
        if line.startswith("## "):
            document.add_heading(line[3:], level=1)
            index += 1
            continue
        if line.startswith("### "):
            document.add_heading(line[4:], level=2)
            index += 1
            continue
        if line.startswith("#### "):
            document.add_heading(line[5:], level=3)
            index += 1
            continue
        numbered = re.match(r"^(\d+)\.\s+(.*)$", line)
        bullet = re.match(r"^[-*]\s+(.*)$", line)
        if numbered:
            paragraph = document.add_paragraph(style="List Number")
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            add_inline_markdown(paragraph, numbered.group(2))
        elif bullet:
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            add_inline_markdown(paragraph, bullet.group(1))
        else:
            paragraph = document.add_paragraph()
            if line.startswith("**Figure "):
                paragraph.style = document.styles["Caption"]
            add_inline_markdown(paragraph, line)
            if "`" in line or "http" in line or re.search(r"[0-9a-f]{40,}", line):
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if line.startswith("**References"):
                paragraph.paragraph_format.keep_with_next = True
            if include_figures:
                plain = line.replace("**", "")
                for prefix, filename in FIGURE_FILES.items():
                    if plain.startswith(prefix):
                        image_paragraph = document.add_paragraph()
                        image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        image_paragraph.paragraph_format.keep_together = True
                        image_paragraph.add_run().add_picture(str(FIGURES / filename), width=Inches(6.2))
                        break
        index += 1


def add_title_page(document: Document) -> None:
    paragraph = document.add_paragraph(style="Title")
    paragraph.paragraph_format.space_before = Pt(54)
    paragraph.add_run("Observed provenance composition is not sampling effort in presence-only species distribution models")

    author = document.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author.paragraph_format.space_before = Pt(26)
    run = author.add_run("Arzaan Ul Mairaj")
    run.bold = True
    run.font.size = Pt(12)

    affiliation = document.add_paragraph()
    affiliation.alignment = WD_ALIGN_PARAGRAPH.CENTER
    affiliation.add_run("Independent researcher, Birmingham, B1 1BA, United Kingdom")

    corresponding = document.add_paragraph()
    corresponding.alignment = WD_ALIGN_PARAGRAPH.CENTER
    corresponding.add_run("Corresponding author: Arzaan Ul Mairaj\narzaaan789@gmail.com")

    target = document.add_paragraph()
    target.alignment = WD_ALIGN_PARAGRAPH.CENTER
    target.paragraph_format.space_before = Pt(20)
    run = target.add_run("Original Research Article prepared for Ecological Modelling")
    run.italic = True

    document.add_page_break()


def build_main() -> Path:
    document = Document()
    configure_document(document)
    add_title_page(document)
    lines = (MANUSCRIPT / "manuscript.md").read_text(encoding="utf-8").splitlines()
    # Skip source title and contact block already represented by the title page.
    abstract_start = lines.index("## Abstract")
    add_markdown(document, lines=lines[abstract_start:], include_figures=True)
    output = MANUSCRIPT / "Ecological_Modelling_manuscript.docx"
    document.save(output)
    return output


def build_supplement() -> Path:
    document = Document()
    configure_document(document)
    add_markdown(document, MANUSCRIPT / "supplement.md")
    output = MANUSCRIPT / "Ecological_Modelling_supplement.docx"
    document.save(output)
    return output


def build_highlights() -> Path:
    document = Document()
    configure_document(document)
    document.add_paragraph("Highlights", style="Title")
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Observed provenance composition is not sampling effort in presence-only species distribution models").italic = True
    document.add_paragraph()
    for line in (MANUSCRIPT / "highlights.txt").read_text(encoding="utf-8").splitlines():
        if line.strip():
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.paragraph_format.space_after = Pt(10)
            paragraph.add_run(line.strip())
    output = MANUSCRIPT / "Ecological_Modelling_highlights.docx"
    document.save(output)
    return output


def build_cover_letter() -> Path:
    document = Document()
    configure_document(document, letter=True)
    lines = (MANUSCRIPT / "cover_letter.md").read_text(encoding="utf-8").splitlines()
    for index, line in enumerate(lines):
        if not line.strip():
            continue
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.line_spacing = 1.15
        paragraph.paragraph_format.space_after = Pt(8)
        add_inline_markdown(paragraph, line)
        if index == 0:
            paragraph.paragraph_format.space_after = Pt(18)
    output = MANUSCRIPT / "Ecological_Modelling_cover_letter.docx"
    document.save(output)
    return output


def main() -> None:
    for path in (build_main(), build_supplement(), build_highlights(), build_cover_letter()):
        print(path)


if __name__ == "__main__":
    main()
